"""Generador documental (specs/doc-generator.md). Propietat per usuari."""

import io
import json
import re as _re
from collections.abc import AsyncIterator
from typing import Annotated, Any, Literal

from fastapi import APIRouter, Depends, Query, UploadFile
from fastapi.responses import Response, StreamingResponse
from pydantic import BaseModel, Field
from sqlalchemy import text
from sqlalchemy.ext.asyncio import AsyncSession

from app.ai import doc_agent, providers
from app.core import authz
from app.core.db import get_session
from app.core.problems import Problem
from app.modules.audit.models import AuditActorType
from app.modules.audit.service import record_audit
from app.modules.users.dependencies import get_request_context
from app.modules.users.service import RequestContext

router = APIRouter(tags=["docgen"])

SessionDep = Annotated[AsyncSession, Depends(get_session)]
ContextDep = Annotated[RequestContext, Depends(get_request_context)]
UseDep = Annotated[authz.AuthzContext, Depends(authz.Authorize("tools:use"))]

DocType = Literal["PPT", "PPA", "REPORT"]


class ProjectBody(BaseModel):
    name: str = Field(min_length=1, max_length=200)


class ReferencesBody(BaseModel):
    document_ids: list[int] = Field(max_length=20)


class ExternalRefBody(BaseModel):
    title: str = Field(min_length=1, max_length=500)
    source_url: str = Field(min_length=10, max_length=1000)
    file_code: str | None = Field(default=None, max_length=100)


class SectionsBody(BaseModel):
    sections: list[dict[str, Any]] = Field(max_length=30)


class DraftBody(BaseModel):
    instructions: str | None = Field(default=None, max_length=2000)
    fields: list[dict[str, Any]] | None = Field(default=None, max_length=10)
    # "draft" redacta de zero; "improve" millora el text manual existent.
    mode: Literal["draft", "improve"] = "draft"
    # Text actual de l'editor (mode improve): evita millorar una versió
    # desactualitzada si el tècnic encara no havia desat.
    content: str | None = Field(default=None, max_length=50000)


async def _own_project(session: AsyncSession, project_id: int, user_id: int) -> dict[str, Any]:
    row = (
        await session.execute(
            text("SELECT id, name, reference_doc_ids, user_id FROM doc_projects WHERE id = :id"),
            {"id": project_id},
        )
    ).first()
    if row is None or row.user_id != user_id:
        raise Problem(404, "Projecte desconegut", "not-found")
    return {"id": row.id, "name": row.name, "reference_doc_ids": row.reference_doc_ids}


async def _audit(
    session: AsyncSession, user_id: int, action: str, resource: str, ctx: RequestContext
) -> None:
    await record_audit(
        session,
        actor_type=AuditActorType.USER,
        action=action,
        success=True,
        actor_id=user_id,
        resource_type="doc_project",
        resource_id=resource,
        ip=ctx.ip,
        user_agent=ctx.user_agent,
        trace_id=ctx.trace_id,
    )


async def _references_detail(session: AsyncSession, ids: list[int]) -> list[dict[str, Any]]:
    if not ids:
        return []
    rows = (
        await session.execute(
            text(
                "SELECT pd.id, pd.title, pd.doc_type, c.file_code FROM phase_documents pd "
                "LEFT JOIN contracts c ON c.id = pd.contract_id WHERE pd.id = ANY(:ids)"
            ),
            {"ids": ids},
        )
    ).all()
    by_id = {r.id: r for r in rows}
    return [
        {
            "id": i,
            "title": by_id[i].title,
            "doc_type": by_id[i].doc_type,
            "file_code": by_id[i].file_code,
        }
        for i in ids
        if i in by_id
    ]


@router.get("/doc-references", operation_id="searchDocReferences")
async def search_doc_references(
    session: SessionDep,
    _authz: UseDep,
    q: Annotated[str | None, Query(min_length=2, max_length=200)] = None,
) -> dict[str, list[dict[str, Any]]]:
    """Documents indexats candidats a referència (per títol o expedient)."""
    conditions = (
        "pd.indexed_at IS NOT NULL AND EXISTS "
        "(SELECT 1 FROM rag_chunks rc WHERE rc.document_id = pd.id)"
    )
    params: dict[str, Any] = {}
    if q:
        escaped = q.replace("\\", "\\\\").replace("%", "\\%").replace("_", "\\_")
        conditions += " AND (pd.title ILIKE :q ESCAPE '\\' OR c.file_code ILIKE :q ESCAPE '\\')"
        params["q"] = f"%{escaped}%"
    rows = (
        await session.execute(
            text(
                "SELECT pd.id, pd.title, pd.doc_type, c.file_code "  # noqa: S608
                "FROM phase_documents pd LEFT JOIN contracts c ON c.id = pd.contract_id "
                f"WHERE {conditions} ORDER BY pd.id DESC LIMIT 20"
            ),
            params,
        )
    ).all()
    return {
        "data": [
            {"id": r.id, "title": r.title, "doc_type": r.doc_type, "file_code": r.file_code}
            for r in rows
        ]
    }


@router.get("/doc-projects", operation_id="listDocProjects")
async def list_projects(session: SessionDep, authz_ctx: UseDep) -> dict[str, list[dict[str, Any]]]:
    rows = (
        await session.execute(
            text(
                "SELECT id, name, reference_doc_ids, created_at FROM doc_projects "
                "WHERE user_id = :u ORDER BY id DESC"
            ),
            {"u": authz_ctx.user.id},
        )
    ).all()
    return {
        "data": [
            {
                "id": r.id,
                "name": r.name,
                "references": len(r.reference_doc_ids or []),
                "created_at": r.created_at,
            }
            for r in rows
        ]
    }


@router.post("/doc-projects", operation_id="createDocProject", status_code=201)
async def create_project(
    body: ProjectBody, session: SessionDep, authz_ctx: UseDep, ctx: ContextDep
) -> dict[str, Any]:
    project_id = (
        await session.execute(
            text("INSERT INTO doc_projects (user_id, name) VALUES (:u, :n) RETURNING id"),
            {"u": authz_ctx.user.id, "n": body.name},
        )
    ).scalar_one()
    for doc_type in ("PPT", "PPA", "REPORT"):
        await session.execute(
            text("INSERT INTO doc_documents (project_id, doc_type) VALUES (:p, :t)"),
            {"p": project_id, "t": doc_type},
        )
    await _audit(session, authz_ctx.user.id, "docgen.project_created", str(project_id), ctx)
    await session.commit()
    return {"id": project_id, "name": body.name}


@router.get("/doc-projects/{id}", operation_id="getDocProject")
async def get_project(id: int, session: SessionDep, authz_ctx: UseDep) -> dict[str, Any]:
    project = await _own_project(session, id, authz_ctx.user.id)
    documents = (
        await session.execute(
            text(
                "SELECT doc_type, sections FROM doc_documents WHERE project_id = :p "
                "ORDER BY doc_type"
            ),
            {"p": id},
        )
    ).all()
    externals = (
        await session.execute(
            text(
                "SELECT id, file_code, title, status, error_detail, chunks_count, expires_at "
                "FROM project_documents WHERE project_id = :p ORDER BY id"
            ),
            {"p": id},
        )
    ).all()
    return {
        "id": project["id"],
        "name": project["name"],
        "references": await _references_detail(session, project["reference_doc_ids"] or []),
        "external_references": [
            {
                "id": r.id,
                "file_code": r.file_code,
                "title": r.title,
                "status": r.status,
                "error_detail": r.error_detail,
                "chunks_count": r.chunks_count,
                "expires_at": r.expires_at,
            }
            for r in externals
        ],
        "documents": {r.doc_type: r.sections for r in documents},
    }


@router.delete("/doc-projects/{id}", operation_id="deleteDocProject", status_code=204)
async def delete_project(id: int, session: SessionDep, authz_ctx: UseDep, ctx: ContextDep) -> None:
    await _own_project(session, id, authz_ctx.user.id)
    await session.execute(text("DELETE FROM doc_projects WHERE id = :id"), {"id": id})
    await _audit(session, authz_ctx.user.id, "docgen.project_deleted", str(id), ctx)
    await session.commit()


@router.put("/doc-projects/{id}/references", operation_id="setDocProjectReferences")
async def set_references(
    id: int, body: ReferencesBody, session: SessionDep, authz_ctx: UseDep, ctx: ContextDep
) -> dict[str, Any]:
    await _own_project(session, id, authz_ctx.user.id)
    valid = list(
        (
            await session.execute(
                text(
                    "SELECT id FROM phase_documents WHERE id = ANY(:ids) AND indexed_at IS NOT NULL"
                ),
                {"ids": body.document_ids},
            )
        ).scalars()
    )
    await session.execute(
        text("UPDATE doc_projects SET reference_doc_ids = CAST(:r AS jsonb) WHERE id = :id"),
        {"r": json.dumps(valid), "id": id},
    )
    await _audit(session, authz_ctx.user.id, "docgen.references_set", str(id), ctx)
    await session.commit()
    return {"references": await _references_detail(session, valid)}


async def _get_sections(session: AsyncSession, project_id: int, doc_type: str) -> list[Any]:
    sections: list[Any] = (
        await session.execute(
            text("SELECT sections FROM doc_documents WHERE project_id = :p AND doc_type = :t"),
            {"p": project_id, "t": doc_type},
        )
    ).scalar_one()
    return sections


async def _save_sections(
    session: AsyncSession, project_id: int, doc_type: str, sections: list[Any]
) -> None:
    await session.execute(
        text(
            "UPDATE doc_documents SET sections = CAST(:s AS jsonb), updated_at = now() "
            "WHERE project_id = :p AND doc_type = :t"
        ),
        {"s": json.dumps(sections), "p": project_id, "t": doc_type},
    )


@router.post(
    "/doc-projects/{id}/external-references",
    operation_id="addExternalReference",
    status_code=202,
)
async def add_external_reference(
    id: int, body: ExternalRefBody, session: SessionDep, authz_ctx: UseDep, ctx: ContextDep
) -> dict[str, Any]:
    """Document del SuperBuscador → índex TEMPORAL d'àmbit del projecte.

    La descàrrega i la indexació van a la cua (crida externa); l'anti-SSRF
    el fa el connector pscp en descarregar.
    """
    await _own_project(session, id, authz_ctx.user.id)
    from app.ai.project_refs import default_expiry
    from app.jobs.service import enqueue_job

    ref_id = (
        await session.execute(
            text(
                "INSERT INTO project_documents (project_id, file_code, title, source_url, "
                "expires_at) VALUES (:p, :f, :t, :u, :e) RETURNING id"
            ),
            {
                "p": id,
                "f": body.file_code,
                "t": body.title,
                "u": body.source_url,
                "e": default_expiry(),
            },
        )
    ).scalar_one()
    job = await enqueue_job(
        session,
        job_type="docgen.index_external",
        payload={"project_document_id": ref_id},
        created_by=authz_ctx.user.id,
        dedup_key=f"docgen.index_external:{ref_id}",
    )
    await _audit(session, authz_ctx.user.id, "docgen.external_ref_added", f"{id}/{ref_id}", ctx)
    await session.commit()
    return {"id": ref_id, "job_id": str(job.id), "status": "pending"}


_UPLOAD_MAX_BYTES = 15 * 1024 * 1024  # PDFs de plecs reals rarament passen de 15 MB


@router.post(
    "/doc-projects/{id}/external-references/upload",
    operation_id="uploadProjectDocument",
    status_code=202,
)
async def upload_project_document(
    id: int,
    session: SessionDep,
    authz_ctx: UseDep,
    ctx: ContextDep,
    file: UploadFile,
) -> dict[str, Any]:
    """PDF propi de l'ordinador de l'usuari → índex TEMPORAL del projecte.

    Mateixa taula i caducitat que les referències del SuperBuscador; el job
    salta la descàrrega perquè el fitxer ja és a l'storage.
    """
    await _own_project(session, id, authz_ctx.user.id)
    filename = (file.filename or "document.pdf").strip()
    if not filename.lower().endswith(".pdf"):
        raise Problem(422, "Només s'accepten fitxers PDF", "validation")
    content = await file.read()
    if len(content) == 0:
        raise Problem(422, "El fitxer és buit", "validation")
    if len(content) > _UPLOAD_MAX_BYTES:
        raise Problem(422, "El fitxer supera el límit de 15 MB", "validation")
    if not content.startswith(b"%PDF"):
        raise Problem(422, "El fitxer no sembla un PDF vàlid", "validation")

    import uuid as _uuid

    from app.ai.project_refs import default_expiry
    from app.core.storage import get_storage
    from app.jobs.service import enqueue_job

    storage_key = f"projects/{id}/{_uuid.uuid4().hex}.pdf"
    await get_storage().put(storage_key, content, "application/pdf")
    ref_id = (
        await session.execute(
            text(
                "INSERT INTO project_documents (project_id, title, storage_key, expires_at) "
                "VALUES (:p, :t, :k, :e) RETURNING id"
            ),
            {"p": id, "t": filename[:500], "k": storage_key, "e": default_expiry()},
        )
    ).scalar_one()
    job = await enqueue_job(
        session,
        job_type="docgen.index_external",
        payload={"project_document_id": ref_id},
        created_by=authz_ctx.user.id,
        dedup_key=f"docgen.index_external:{ref_id}",
    )
    await _audit(session, authz_ctx.user.id, "docgen.document_uploaded", f"{id}/{ref_id}", ctx)
    await session.commit()
    return {"id": ref_id, "job_id": str(job.id), "status": "pending"}


@router.delete(
    "/doc-projects/{id}/external-references/{ref_id}",
    operation_id="removeExternalReference",
    status_code=204,
)
async def remove_external_reference(
    id: int, ref_id: int, session: SessionDep, authz_ctx: UseDep, ctx: ContextDep
) -> None:
    await _own_project(session, id, authz_ctx.user.id)
    await session.execute(
        text("DELETE FROM project_documents WHERE id = :r AND project_id = :p"),
        {"r": ref_id, "p": id},
    )
    await _audit(session, authz_ctx.user.id, "docgen.external_ref_removed", f"{id}/{ref_id}", ctx)
    await session.commit()


@router.patch("/doc-projects/{id}/documents/{doc_type}", operation_id="saveDocSections")
async def save_doc_sections(
    id: int,
    doc_type: DocType,
    body: SectionsBody,
    session: SessionDep,
    authz_ctx: UseDep,
    ctx: ContextDep,
) -> dict[str, Any]:
    await _own_project(session, id, authz_ctx.user.id)
    cleaned = [
        {
            "title": str(s.get("title", ""))[:300],
            "instructions": str(s.get("instructions", ""))[:2000],
            "content_md": str(s.get("content_md", ""))[:50000],
            "sources": s.get("sources") or [],
            "fields": [
                {
                    "label": str(f.get("label", ""))[:120],
                    "hint": str(f.get("hint", ""))[:300],
                    "value": str(f.get("value", ""))[:500],
                }
                for f in (s.get("fields") or [])
                if isinstance(f, dict)
            ][:10],
        }
        for s in body.sections
    ]
    await _save_sections(session, id, doc_type, cleaned)
    await _audit(session, authz_ctx.user.id, "docgen.sections_saved", f"{id}/{doc_type}", ctx)
    await session.commit()
    return {"sections": cleaned}


@router.post(
    "/doc-projects/{id}/documents/{doc_type}/actions/generate-index",
    operation_id="generateDocIndex",
)
async def generate_doc_index(
    id: int, doc_type: DocType, session: SessionDep, authz_ctx: UseDep, ctx: ContextDep
) -> dict[str, Any]:
    await _own_project(session, id, authz_ctx.user.id)
    try:
        titles = await doc_agent.generate_index(
            session, id, doc_type, user_id=authz_ctx.user.id, trace_id=ctx.trace_id
        )
    except providers.ProviderError as exc:
        raise Problem(502, "El proveïdor d'IA no ha respost", "upstream", detail=str(exc)) from None
    sections = [
        {
            "title": t["title"],
            "instructions": "",
            "content_md": "",
            "sources": [],
            "fields": t.get("fields", []),
        }
        for t in titles
    ]
    await _save_sections(session, id, doc_type, sections)
    await _audit(session, authz_ctx.user.id, "docgen.index_generated", f"{id}/{doc_type}", ctx)
    await session.commit()
    return {"sections": sections}


@router.post(
    "/doc-projects/{id}/documents/{doc_type}/sections/{section_index}/actions/draft/stream",
    operation_id="draftDocSectionStream",
)
async def draft_section_stream(
    id: int,
    doc_type: DocType,
    section_index: int,
    body: DraftBody,
    session: SessionDep,
    authz_ctx: UseDep,
    ctx: ContextDep,
) -> StreamingResponse:
    await _own_project(session, id, authz_ctx.user.id)
    sections = await _get_sections(session, id, doc_type)
    if not (0 <= section_index < len(sections)):
        raise Problem(404, "Secció desconeguda", "not-found")
    section = sections[section_index]
    if body.instructions is not None:
        section["instructions"] = body.instructions[:2000]

    def _line(payload: dict[str, Any]) -> str:
        return json.dumps(payload, ensure_ascii=False, default=str) + "\n"

    async def generate() -> AsyncIterator[str]:
        collected: list[str] = []
        sources: list[dict[str, Any]] = []
        try:
            draft_fields = body.fields if body.fields is not None else section.get("fields") or []
            improve_text = None
            if body.mode == "improve":
                improve_text = body.content or str(section.get("content_md", ""))
            async for event in doc_agent.draft_section_events(
                session,
                id,
                doc_type,
                str(section.get("title", "")),
                body.instructions or str(section.get("instructions", "")) or None,
                fields=draft_fields,
                improve_text=improve_text or None,
                user_id=authz_ctx.user.id,
                trace_id=ctx.trace_id,
            ):
                if event["type"] == "delta":
                    collected.append(str(event["text"]))
                elif event["type"] == "sources":
                    sources = event["sources"]
                yield _line(event)
            # Desat ATÒMIC de només aquesta secció (jsonb_set sobre l'estat
            # actual de la BD): redaccions concurrents d'altres seccions no
            # es trepitgen (el desat de l'array sencer feia last-writer-wins).
            await session.execute(
                text(
                    "UPDATE doc_documents SET sections = jsonb_set(jsonb_set(jsonb_set("
                    "jsonb_set(sections, ARRAY[:i_txt, 'content_md'], "
                    "to_jsonb(CAST(:content AS text))), "
                    "ARRAY[:i_txt, 'sources'], CAST(:sources AS jsonb)), "
                    "ARRAY[:i_txt, 'instructions'], to_jsonb(CAST(:instructions AS text))), "
                    "ARRAY[:i_txt, 'fields'], CAST(:fields AS jsonb)), "
                    "updated_at = now() "
                    "WHERE project_id = :p AND doc_type = :t "
                    "AND jsonb_array_length(sections) > :i_int"
                ),
                {
                    "i_txt": str(section_index),
                    "i_int": section_index,
                    "content": "".join(collected)[:50000],
                    "sources": json.dumps(sources),
                    "instructions": str(
                        body.instructions
                        if body.instructions is not None
                        else section.get("instructions", "")
                    )[:2000],
                    "fields": json.dumps(draft_fields or []),
                    "p": id,
                    "t": doc_type,
                },
            )
            await _audit(
                session,
                authz_ctx.user.id,
                "docgen.section_drafted",
                f"{id}/{doc_type}/{section_index}",
                ctx,
            )
            await session.commit()
            yield _line({"type": "done"})
        except providers.ProviderError as exc:
            yield _line({"type": "error", "detail": str(exc)})

    return StreamingResponse(
        generate(),
        media_type="application/x-ndjson",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


def _md_to_docx(document: Any, markdown: str) -> None:
    """Conversió línia a línia: NOMÉS una línia que comença per # és títol
    (mai el paràgraf que la segueix); llistes -, * i numerades; la resta,
    paràgrafs acumulats fins a línia en blanc."""
    paragraph_lines: list[str] = []

    def flush() -> None:
        if paragraph_lines:
            document.add_paragraph(" ".join(paragraph_lines).replace("**", ""))
            paragraph_lines.clear()

    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line:
            flush()
            continue
        heading = _re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            flush()
            level = min(2 + len(heading.group(1)) - 1, 4)
            document.add_heading(heading.group(2).replace("**", "").strip(), level=level)
            continue
        if _re.match(r"^[-*]\s+", line):
            flush()
            document.add_paragraph(
                _re.sub(r"^[-*]\s+", "", line).replace("**", ""), style="List Bullet"
            )
            continue
        if _re.match(r"^\d+[.)]\s+", line):
            flush()
            document.add_paragraph(
                _re.sub(r"^\d+[.)]\s+", "", line).replace("**", ""), style="List Number"
            )
            continue
        paragraph_lines.append(line)
    flush()


@router.post(
    "/doc-projects/{id}/documents/{doc_type}/actions/review/stream",
    operation_id="reviewDocStream",
)
async def review_document_stream(
    id: int, doc_type: DocType, session: SessionDep, authz_ctx: UseDep, ctx: ContextDep
) -> StreamingResponse:
    """Agent revisor: segona opinió sobre el document sencer (mai reescriu)."""
    await _own_project(session, id, authz_ctx.user.id)
    sections = await _get_sections(session, id, doc_type)

    def line(payload: dict[str, Any]) -> str:
        return json.dumps(payload, ensure_ascii=False, default=str) + "\n"

    async def generate() -> AsyncIterator[str]:
        try:
            async for event in doc_agent.review_document_events(
                session,
                doc_type,
                sections,
                user_id=authz_ctx.user.id,
                trace_id=ctx.trace_id,
            ):
                yield line(event)
            yield line({"type": "done"})
        except providers.ProviderError as exc:
            yield line({"type": "error", "detail": str(exc)})

    return StreamingResponse(
        generate(),
        media_type="application/x-ndjson",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@router.get("/doc-projects/{id}/documents/{doc_type}/export.docx", operation_id="exportDocDocx")
async def export_docx(
    id: int, doc_type: DocType, session: SessionDep, authz_ctx: UseDep, ctx: ContextDep
) -> Response:
    project = await _own_project(session, id, authz_ctx.user.id)
    sections = await _get_sections(session, id, doc_type)

    from docx import Document

    document = Document()
    document.add_heading(doc_agent.DOC_TYPE_NAMES[doc_type], level=0)
    document.add_paragraph(f"Projecte: {project['name']} — Ajuntament de Cunit")
    all_sources: dict[str, str] = {}
    for section in sections:
        document.add_heading(str(section.get("title", "")), level=1)
        for block in str(section.get("content_md", "")).split("\n\n"):
            line = block.strip()
            if not line:
                continue
            if line.startswith("#"):
                document.add_heading(line.lstrip("# ").strip(), level=2)
            elif line.startswith(("- ", "* ")):
                for item in line.splitlines():
                    document.add_paragraph(item.lstrip("-* ").strip(), style="List Bullet")
            else:
                document.add_paragraph(line.replace("**", ""))
        for source in section.get("sources") or []:
            key = f"{source.get('file_code')} — {source.get('document_title')}"
            all_sources[key] = key
    if all_sources:
        document.add_heading("Fonts consultades", level=1)
        for key in all_sources:
            document.add_paragraph(key, style="List Bullet")

    buffer = io.BytesIO()
    document.save(buffer)
    await _audit(session, authz_ctx.user.id, "docgen.exported", f"{id}/{doc_type}", ctx)
    await session.commit()
    filename = f"{doc_type}-{project['name'][:40]}.docx".replace(" ", "-")
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
